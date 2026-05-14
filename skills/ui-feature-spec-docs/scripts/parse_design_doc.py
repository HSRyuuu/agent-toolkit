#!/usr/bin/env python3
"""화면정의서(docx/pdf) 파싱 — 텍스트·표를 구조화 JSON으로 출력.

LLM이 직접 읽고 화면 단위로 분리·매핑할 수 있도록 raw에 가까운 형태로 추출한다.
의미 해석(어느 화면인지, 어느 기능인지)은 호출자가 담당.

사용:
    python3 parse_design_doc.py <input.docx|.pdf> [--output OUT.json]
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any


def parse_docx(path: Path) -> dict[str, Any]:
    try:
        from docx import Document
    except ImportError:
        sys.stderr.write(
            "python-docx가 설치되어 있지 않습니다. "
            "`pip install -r scripts/requirements.txt` 실행 후 재시도하세요.\n"
        )
        sys.exit(2)

    doc = Document(str(path))

    paragraphs: list[dict[str, str]] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        paragraphs.append(
            {
                "style": p.style.name if p.style else "Normal",
                "text": text,
            }
        )

    tables: list[list[list[str]]] = []
    for t in doc.tables:
        rows: list[list[str]] = []
        for row in t.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)

    return {
        "type": "docx",
        "source": str(path),
        "paragraphs": paragraphs,
        "tables": tables,
    }


def parse_pdf(path: Path) -> dict[str, Any]:
    try:
        from pypdf import PdfReader
    except ImportError:
        sys.stderr.write(
            "pypdf가 설치되어 있지 않습니다. "
            "`pip install -r scripts/requirements.txt` 실행 후 재시도하세요.\n"
        )
        sys.exit(2)

    reader = PdfReader(str(path))
    pages: list[dict[str, Any]] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        pages.append(
            {
                "page": i + 1,
                "text": text.strip(),
            }
        )

    return {
        "type": "pdf",
        "source": str(path),
        "pages": pages,
    }


def main() -> None:
    parser = argparse.ArgumentParser(
        description="화면정의서(docx/pdf) 파싱 → JSON",
    )
    parser.add_argument("input", type=Path, help="입력 docx 또는 pdf 경로")
    parser.add_argument(
        "--output",
        "-o",
        type=Path,
        default=None,
        help="출력 JSON 경로 (생략 시 stdout)",
    )
    args = parser.parse_args()

    if not args.input.exists():
        sys.stderr.write(f"파일을 찾을 수 없습니다: {args.input}\n")
        sys.exit(1)

    suffix = args.input.suffix.lower()
    if suffix == ".docx":
        result = parse_docx(args.input)
    elif suffix == ".pdf":
        result = parse_pdf(args.input)
    else:
        sys.stderr.write(
            f"지원하지 않는 확장자: {suffix} (docx 또는 pdf만 지원)\n"
        )
        sys.exit(1)

    out_text = json.dumps(result, ensure_ascii=False, indent=2)

    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(out_text, encoding="utf-8")
        print(f"Wrote {args.output}")
    else:
        print(out_text)


if __name__ == "__main__":
    main()
